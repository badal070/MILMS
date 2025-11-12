import re
from docx import Document
from typing import List, Dict, Optional, Tuple
import PyPDF2
from io import BytesIO
import os 

class QuestionParser:
    """Robust parser for quiz questions from Word/PDF files"""
    
    def __init__(self):
        self.question_endings = ['?', '.', ':']
        self.correct_marker = '*'
        
    def parse_from_docx(self, file_path: str) -> List[Dict]:
        """Parse questions from Word document"""
        try:
            doc = Document(file_path)
            text_content = self._extract_text_from_docx(doc)
            return self._parse_questions(text_content)
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def parse_from_pdf(self, file_path: str) -> List[Dict]:
        """Parse questions from PDF document"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                
                combined_text = '\n'.join(text_content)
                return self._parse_questions(combined_text)
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    def _extract_text_from_docx(self, doc: Document) -> str:
        """Extract clean text from Word document"""
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        
        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)
        
        return '\n'.join(lines)
    
    def _parse_questions(self, text: str) -> List[Dict]:
        """
        Core parsing logic - handles messy structures
        Returns list of question dictionaries
        """
        questions = []
        lines = text.split('\n')
        
        # Clean and normalize lines
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            # Remove numbering like "1.", "Q1.", "Question 1:"
            line = re.sub(r'^(\d+\.?\s*|\(?[a-zA-Z]\)\.?\s*|Q\d+\.?\s*|Question\s+\d+\.?\s*)', '', line)
            if line:
                cleaned_lines.append(line)
        
        i = 0
        while i < len(cleaned_lines):
            # Detect question
            if self._is_question(cleaned_lines[i]):
                question_text = cleaned_lines[i]
                i += 1
                
                # Collect options
                options = []
                while i < len(cleaned_lines) and len(options) < 4:
                    line = cleaned_lines[i]
                    
                    # Stop if next question detected
                    if self._is_question(line) and len(options) > 0:
                        break
                    
                    # Check if it's an option
                    if self._is_option(line):
                        option_text, is_correct = self._parse_option(line)
                        if option_text:
                            options.append({
                                'text': option_text,
                                'is_correct': is_correct
                            })
                    
                    i += 1
                
                # Validate and add question
                if len(options) == 4 and sum(1 for opt in options if opt['is_correct']) == 1:
                    questions.append({
                        'question': question_text,
                        'options': options
                    })
                elif len(options) > 0:
                    # Try to fix common issues
                    fixed_question = self._fix_question_structure(question_text, options)
                    if fixed_question:
                        questions.append(fixed_question)
            else:
                i += 1
        
        return questions
    
    def _is_question(self, text: str) -> bool:
        """Check if text is a question"""
        if not text or len(text) < 5:
            return False
        
        # Check if ends with question markers
        ends_with_marker = any(text.endswith(ending) for ending in self.question_endings)
        
        # Check if contains question words
        question_words = ['what', 'which', 'where', 'when', 'who', 'whom', 'whose', 
                         'why', 'how', 'is', 'are', 'was', 'were', 'do', 'does', 
                         'did', 'can', 'could', 'will', 'would', 'should']
        
        first_word = text.split()[0].lower() if text.split() else ''
        starts_with_question = first_word in question_words
        
        # Must end with marker OR start with question word and be long enough
        return ends_with_marker or (starts_with_question and len(text) > 10)
    
    def _is_option(self, text: str) -> bool:
        """Check if text is an option"""
        if not text or len(text) < 1:
            return False
        
        # Check for option patterns
        # Matches: "a)", "A.", "a-", "a ", "(a)", "[a]", etc.
        option_pattern = r'^[\(\[]?[a-dA-D][\)\]\.:\-\s]'
        
        # Also accept lines that don't look like questions
        is_short = len(text) < 100
        not_question = not self._is_question(text)
        
        return bool(re.match(option_pattern, text)) or (is_short and not_question)
    
    def _parse_option(self, text: str) -> Tuple[str, bool]:
        """
        Parse option text and check if correct
        Returns (option_text, is_correct)
        """
        # Check if marked as correct
        is_correct = text.rstrip().endswith(self.correct_marker)
        
        # Remove correct marker
        if is_correct:
            text = text.rstrip()[:-1].strip()
        
        # Remove option prefix (a), A., etc.)
        text = re.sub(r'^[\(\[]?[a-dA-D][\)\]\.:\-\s]+', '', text).strip()
        
        return text, is_correct
    
    def _fix_question_structure(self, question_text: str, options: List[Dict]) -> Optional[Dict]:
        """
        Try to fix questions with incorrect structure
        """
        # If we have exactly one correct answer but wrong number of options
        correct_count = sum(1 for opt in options if opt['is_correct'])
        
        if correct_count != 1:
            # Try to find or assign correct answer
            if correct_count == 0 and len(options) >= 4:
                # Mark first option as correct as fallback
                options[0]['is_correct'] = True
                correct_count = 1
            elif correct_count > 1:
                # Keep only first correct answer
                found_correct = False
                for opt in options:
                    if opt['is_correct'] and not found_correct:
                        found_correct = True
                    else:
                        opt['is_correct'] = False
                correct_count = 1
        
        # If we have wrong number of options
        if len(options) < 4:
            # Add placeholder options
            while len(options) < 4:
                options.append({
                    'text': f'Option {len(options) + 1}',
                    'is_correct': False
                })
        elif len(options) > 4:
            # Keep first 4 options (ensure correct answer is included)
            correct_options = [opt for opt in options if opt['is_correct']]
            other_options = [opt for opt in options if not opt['is_correct']]
            
            if correct_options:
                options = correct_options[:1] + other_options[:3]
            else:
                options = options[:4]
        
        if len(options) == 4 and correct_count == 1:
            return {
                'question': question_text,
                'options': options
            }
        
        return None


# Main parsing functions for use in Django
def parse_question_from_docx(file_path: str) -> List[Dict]:
    """Parse questions from Word document with error handling"""
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        parser = QuestionParser()
        questions = parser.parse_from_docx(file_path)
        
        if not questions:
            raise ValueError("No valid questions found in the document. Please check the format.")
        
        return questions
    except Exception as e:
        raise Exception(f"Error parsing Word document: {str(e)}")


def parse_question_from_pdf(file_path: str) -> List[Dict]:
    """Parse questions from PDF with error handling"""
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        parser = QuestionParser()
        questions = parser.parse_from_pdf(file_path)
        
        if not questions:
            raise ValueError("No valid questions found in the PDF. Please check the format.")
        
        return questions
    except Exception as e:
        raise Exception(f"Error parsing PDF: {str(e)}")


# Add missing import at top




# Activity logging utility
def log_activity(user, action, description='', request=None):
    """Log user activity with IP address"""
    from .models import ActivityLog
    
    ip_address = None
    if request:
        x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
        if x_forwarded_for:
            ip_address = x_forwarded_for.split(',')[0]
        else:
            ip_address = request.META.get('REMOTE_ADDR')
    
    ActivityLog.objects.create(
        user=user,
        action=action,
        description=description,
        ip_address=ip_address
    )


def get_client_ip(request):
    """Extract client IP address from request"""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        return x_forwarded_for.split(',')[0]
    return request.META.get('REMOTE_ADDR')


# Validation and debugging utilities
def validate_questions(questions: List[Dict]) -> Tuple[bool, List[str]]:
    """
    Validate parsed questions
    
    Returns:
        (is_valid, list_of_errors)
    """
    errors = []
    
    for idx, q in enumerate(questions, 1):
        # Check question text
        if not q.get('question'):
            errors.append(f"Question {idx}: Missing question text")
        
        # Check options
        options = q.get('options', [])
        if len(options) != 4:
            errors.append(f"Question {idx}: Must have exactly 4 options (found {len(options)})")
        
        # Check correct answer
        correct_count = sum(1 for opt in options if opt.get('is_correct'))
        if correct_count != 1:
            errors.append(f"Question {idx}: Must have exactly 1 correct answer (found {correct_count})")
        
        # Check option text
        for opt_idx, opt in enumerate(options, 1):
            if not opt.get('text'):
                errors.append(f"Question {idx}, Option {opt_idx}: Missing option text")
    
    return len(errors) == 0, errors


def preview_parsed_questions(questions: List[Dict], max_questions: int = 3) -> str:
    """
    Generate preview text for parsed questions
    Useful for debugging and user feedback
    """
    preview = []
    
    for idx, q in enumerate(questions[:max_questions], 1):
        preview.append(f"\n{'='*50}")
        preview.append(f"Question {idx}: {q['question']}")
        preview.append(f"{'-'*50}")
        
        for opt_idx, opt in enumerate(q['options'], 1):
            marker = " ✓ [CORRECT]" if opt['is_correct'] else ""
            preview.append(f"  {chr(64+opt_idx)}. {opt['text']}{marker}")
    
    preview.append(f"\n{'='*50}")
    preview.append(f"Total questions parsed: {len(questions)}")
    
    return '\n'.join(preview)

    


def parse_descriptive_questions_from_docx(file_path):
    """
    Parse descriptive questions from Word document
    
    Expected Format:
    Q: Question text here?
    Marks: 10
    Word Limit: 500
    Reference Answer: The answer here...
    Guidelines: Key points to look for...
    ---
    
    Returns list of question dictionaries
    """
    doc = Document(file_path)
    questions = []
    current_question = {}
    current_field = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check for separator
        if text == '---' or text.startswith('==='):
            if current_question.get('question'):
                questions.append(current_question)
                current_question = {}
                current_field = None
            continue
        
        # Check for field markers
        if text.startswith('Q:') or text.startswith('Question:'):
            current_field = 'question'
            current_question['question'] = text.split(':', 1)[1].strip()
        
        elif text.startswith('Marks:'):
            try:
                marks = re.findall(r'\d+', text)[0]
                current_question['max_marks'] = int(marks)
            except:
                current_question['max_marks'] = 10
            current_field = None
        
        elif text.startswith('Word Limit:'):
            try:
                limit = re.findall(r'\d+', text)[0]
                current_question['word_limit'] = int(limit)
            except:
                current_question['word_limit'] = 500
            current_field = None
        
        elif text.startswith('Reference Answer:'):
            current_field = 'reference_answer'
            current_question['reference_answer'] = text.split(':', 1)[1].strip()
        
        elif text.startswith('Guidelines:') or text.startswith('Marking Guidelines:'):
            current_field = 'marking_guidelines'
            current_question['marking_guidelines'] = text.split(':', 1)[1].strip()
        
        else:
            # Continue current field
            if current_field and current_field in current_question:
                current_question[current_field] += '\n' + text
    
    # Add last question
    if current_question.get('question'):
        questions.append(current_question)
    
    # Set defaults
    for q in questions:
        q.setdefault('max_marks', 10)
        q.setdefault('word_limit', 500)
        q.setdefault('reference_answer', '')
        q.setdefault('marking_guidelines', '')
    
    return questions


def format_ai_feedback_html(evaluation_data):
    """Format AI evaluation data into readable HTML"""
    if not evaluation_data:
        return ''
    
    html = '<div class="ai-feedback">'
    
    # Overall score
    html += f'<div class="mb-3">'
    html += f'<h6>Overall Score: {evaluation_data.get("overall_score", 0)}/{evaluation_data.get("max_score", 100)}</h6>'
    html += f'<div class="progress"><div class="progress-bar" style="width: {evaluation_data.get("percentage", 0)}%">{evaluation_data.get("percentage", 0)}%</div></div>'
    html += '</div>'
    
    # Detailed scores
    html += '<div class="row mb-3">'
    scores = [
        ('Spelling', evaluation_data.get('spelling_analysis', {}).get('spelling_score', 0)),
        ('Relevance', evaluation_data.get('relevance_analysis', {}).get('relevance_score', 0)),
        ('Content', evaluation_data.get('content_analysis', {}).get('content_score', 0)),
        ('Grammar', evaluation_data.get('grammar_analysis', {}).get('grammar_score', 0)),
    ]
    
    for label, score in scores:
        html += f'<div class="col-3"><small class="text-muted">{label}</small><br><strong>{score}/10</strong></div>'
    
    html += '</div>'
    
    # Feedback
    feedback = evaluation_data.get('feedback', '')
    if feedback:
        html += f'<div class="alert alert-info"><strong>Feedback:</strong> {feedback}</div>'
    
    html += '</div>'
    return html


def calculate_weighted_score(ai_score, manual_score, ai_weight):
    """Calculate weighted final score"""
    if ai_score is None or manual_score is None:
        return manual_score or ai_score or 0
    
    return (ai_score * ai_weight) + (manual_score * (1 - ai_weight))